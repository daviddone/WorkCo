# -*- coding: utf8 -*-
""" 
    Парсинг лога fiddler и загрузка в БД
    Таблица:
    drop table mk$oapi_requests;
    create table mk$oapi_requests(req_id number(8),request varchar2(128), params varchar2(128));
    select * from mk$oapi_requests
"""

import sys, tempfile, shutil, zipfile, os
from bs4 import BeautifulSoup
import json
import docx

def create_temp_dir():
    try:
        return tempfile.mkdtemp()
    except:
        sys.exit(-1)

def clean_temp_dir(tmpdir):
    if tmpdir: 
        try:
            shutil.rmtree(tmpdir)
        except:
            print("failed to clean up tmpdir %s you will have to do it" % (tmpdir))

def unzip_archive(archive, tmpdir):
    try:
        z = zipfile.ZipFile(archive,"r")
    except:
        sys.exit(-1)
    try:
       z.extractall(tmpdir)
       z.close()
    except:
       sys.exit(-1)


def parse_index_file(tmpdir):
    with open(os.path.join(tmpdir, '_index.htm'), 'r',errors='ignore') as data:
        requests = []
        soup = BeautifulSoup(data)
        docfile = docx.Document()
        for row in soup.findAll('tr')[1:]:
            req_id = int(row.find('td').find('a')['href'].split('\\')[1].split('_')[0])
            td = row.findAll('td')[5].text
            print(td)
            requests.append({'req_id': req_id, 'request': td, 'params': ''})
            req_file = row.findAll('td')[0].findAll('a')[0]['href']
            response_file = row.findAll('td')[0].findAll('a')[1]['href']
            print('Used request file: %s' % req_file)
            print('Used response_file file: %s' % response_file)
            rfile = open(os.path.join(tmpdir, req_file), 'r',encoding="utf-8",errors='ignore')
            resultfile = open(os.path.join(tmpdir, response_file), 'r',encoding="utf-8",errors='ignore')
            rfilelines = rfile.readlines()
            resultlines = resultfile.readlines()
            url = rfilelines[0].split("HTTP")[0]  #访问路径
            params = rfilelines[-1]  #访问路径
            resultline = ""
            for item in resultlines:
                if item.startswith("{\""):
                    resultline = item
            print ("请求路径：%s"%url)
            print ("请求参数：%s"%params)
            print ("请求响应：%s"%resultline)
            model = url.split("/")[-1].strip()
            docfile.add_heading(model, level=2)
            docfile.add_paragraph("请求路径：%s"%url)
            docfile.add_paragraph("请求参数：%s"%params)
            docfile.add_paragraph("请求响应：%s"%resultline)
        docfile.save("策略下发接口.docx")
        return requests


def work_with_archive(archive, tmpdir):
    unzip_archive(archive, tmpdir)
    requests = parse_index_file(tmpdir)

if __name__ == '__main__':
    if (len(sys.argv) > 1):
        print('Start saz parser')
        tmpdir = create_temp_dir()
        work_with_archive(sys.argv[1], tmpdir)
        clean_temp_dir(tmpdir)
        print('Stop saz parser')
    else:
        print(u'Usage ./saz_parser.py {file}')